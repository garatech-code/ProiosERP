from docx import Document

files = [
    "/app/apps/operaciones/templates/Proios_Cotizacion_TEMPLATE_ES.docx",
    "/app/apps/operaciones/templates/Proios_Quotation_TEMPLATE_EN.docx"
]

def extract_tags(doc_path):
    try:
        doc = Document(doc_path)
        text_blocks = []
        for p in doc.paragraphs:
            if '{{' in p.text or '{%' in p.text:
                text_blocks.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '{{' in p.text or '{%' in p.text:
                            text_blocks.append(p.text)
        print(f"--- {doc_path} ---")
        for block in text_blocks:
            print(block.strip())
    except Exception as e:
        print(f"Error reading {doc_path}: {e}")

for f in files:
    extract_tags(f)
