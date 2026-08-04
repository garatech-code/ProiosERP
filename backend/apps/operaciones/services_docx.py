import io
import os
import re
import json
from django.conf import settings
from datetime import datetime
from apps.inventario.models import Articulo
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn

def insert_header_logo(document):
    """
    Busca la primera imagen disponible en la raíz del proyecto y la inserta
    en el encabezado (alineada a la derecha).
    """
    valid_logos = ['logo.png', 'header.png', 'brand.png', 'logo.jpg']
    logo_path = None
    
    for img in valid_logos:
        p = os.path.join(settings.BASE_DIR, 'static_local', img)
        if os.path.exists(p):
            logo_path = p
            break

    if logo_path:
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Cm(5.0))
        
        spacing_paragraph = header.add_paragraph()
        spacing_paragraph.paragraph_format.space_after = Pt(15)

def set_cell_border(cell, **kwargs):
    """
    Función de ayuda para establecer bordes finos en las celdas de las tablas.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def generar_cotizacion_docx(operacion, offer_validity="15 days", payment_terms="30 days from invoice date", delivery_time="5", include_vat=True, scope_includes="[detail what the supply / service comprises]", scope_excludes="[freight, customs clearance, additional labour, parts not listed, etc.]", notes="[Other relevant note]", attn="Operations / Technical Department"):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Tahoma'
    
    # 1. Configurar márgenes de la página (Letter Size)
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.08)
    section.bottom_margin = Cm(2.08)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # 2. HEADER
    insert_header_logo(doc)

    # 3. TÍTULO PRINCIPAL
    title = doc.add_paragraph()
    title_run = title.add_run("QUOTATION")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(20)

    # 4. BLOQUE SUPERIOR (3 COLUMNAS - Usamos una tabla sin bordes)
    top_table = doc.add_table(rows=1, cols=3)
    top_table.autofit = False
    
    # Variables a reemplazar
    cliente = operacion.cliente.name if operacion.cliente else ''
    buque = operacion.ship.name if operacion.ship else ''
    puerto = operacion.port.name if operacion.port else ''
    eta_str = operacion.eta.strftime('%d/%m/%Y %H:%M') if operacion.eta else ''
    ref_num = f"PS-COT-{operacion.fecha_creacion.year}-{operacion.id:03d}" if operacion.fecha_creacion else f"PS-COT-2026-{operacion.id:03d}"
    
    # Columna 1: FROM
    cell_1 = top_table.cell(0, 0)
    cell_1.width = Cm(6.0)
    p1 = cell_1.paragraphs[0]
    p1.add_run("FROM\n").bold = True
    p1.add_run("Proios S.A.\nComodoro Pedro Zanni 351 floor 5th 503 LN.\nBuenos Aires (C1104AAH) Argentina\n\nTel: +549 11 57265031 · eva@proios.com")
    p1.runs[0].font.size = Pt(10)
    
    # Columna 2: TO
    cell_2 = top_table.cell(0, 1)
    cell_2.width = Cm(6.0)
    p2 = cell_2.paragraphs[0]
    p2.add_run("TO\n").bold = True
    p2.add_run(f"{cliente}\n\nAttn:\n{attn}")

    # Columna 3: QUOTATION INFO
    cell_3 = top_table.cell(0, 2)
    cell_3.width = Cm(6.0)
    p3 = cell_3.paragraphs[0]
    p3.add_run("QUOTATION\n").bold = True
    current_date = datetime.now().strftime("%d %b %Y")
    p3.add_run(f"No. {ref_num}\nDate: {current_date}\nValid: {offer_validity}\n")
    p3.add_run(f"Vessel: {buque}\nPort: {puerto}\nETA: {eta_str}")

    doc.add_paragraph() # Espacio

    # Determinar tipo de operación en inglés
    tipo_map = {
        'productos': 'products',
        'quimicos': 'chemicals',
        'servicios': 'services',
        'otros': 'spare parts'
    }
    tipo_en = tipo_map.get(operacion.tipo_operacion, 'products')

    # 5. PÁRRAFO INTRODUCTORIO
    intro = doc.add_paragraph(f"We are pleased to submit our quotation for the supply of {tipo_en}, as detailed below:")
    intro.paragraph_format.space_after = Pt(15)

    # 6. SECCIÓN ITEMS
    items_title = doc.add_paragraph()
    run = items_title.add_run("ITEMS")
    run.bold = True
    run.font.size = Pt(8)

    headers = ["No.", "Description", "Qty", "Unit", "Unit price", "Amount"]
    items_table = doc.add_table(rows=1, cols=len(headers))
    items_table.style = 'Table Grid'
    
    # Llenar encabezados
    hdr_cells = items_table.rows[0].cells
    for i, text in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(text)
        run.bold = True
        
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    def set_cell_bg(cell, color):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
    # Obtener detalles
    total_usd = 0
    detalles = list(operacion.detalles.all())
    if detalles:
        grupos = {}
        for det in detalles:
            try:
                articulo = Articulo.objects.get(id=det.articulo_id)
                cat = str(articulo.categoria).strip() if articulo.categoria and str(articulo.categoria).strip() else "GENERAL"
                desc = articulo.nombre
                unit = str(articulo.unidad) if articulo.unidad else "u"
            except Articulo.DoesNotExist:
                cat = "GENERAL"
                desc = f"Item #{det.articulo_id}"
                unit = "u"
            cat = cat.upper()
            if cat not in grupos:
                grupos[cat] = []
            grupos[cat].append({
                'desc': desc,
                'qty': float(det.cantidad),
                'unit': unit,
                'price': float(det.precio_unitario) if det.precio_unitario else 0.0
            })
            
        idx = 1
        for cat_name, items in grupos.items():
            # Category Row
            cat_row = items_table.add_row()
            a = cat_row.cells[0]
            b = cat_row.cells[5]
            a.merge(b)
            p = a.paragraphs[0]
            r = p.add_run(cat_name)
            r.bold = True
            set_cell_bg(a, "E2E8F0")
            
            subtotal_cat = 0
            cant_cat = 0
            for item in items:
                qty = item['qty']
                price = item['price']
                amount = qty * price
                subtotal_cat += amount
                cant_cat += qty
                total_usd += amount
                
                row = items_table.add_row().cells
                row[0].text = str(idx)
                row[1].text = str(item['desc'])
                row[2].text = str(int(qty) if qty.is_integer() else qty)
                row[3].text = str(item['unit'])
                row[4].text = f"{price:.2f}"
                row[5].text = f"{amount:.2f}"
                idx += 1
                
            # Subtotal Row
            sub_row = items_table.add_row()
            a = sub_row.cells[0]
            b = sub_row.cells[4]
            a.merge(b)
            a.paragraphs[0].add_run(f"Subtotal {cat_name.lower()} ({int(cant_cat)} {items[0]['unit'] if len(items)==1 else 'u.'})").bold = True
            sub_row.cells[5].paragraphs[0].add_run(f"{subtotal_cat:.2f}").bold = True
    else:
        row_cells = items_table.add_row().cells
        example_data = ["1", operacion.detalle_servicio or "[Description]", "0", "u", "0.00", "0.00"]
        for i, text in enumerate(example_data):
            row_cells[i].text = text


    doc.add_paragraph() # Espaciado

    # Totales
    totals_p = doc.add_paragraph()
    totals_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if str(include_vat).lower() == 'true' or include_vat is True:
        totals_p.add_run(f"Subtotal:       {total_usd:.2f}\n")
        totals_p.add_run(f"VAT (21%):      {(float(total_usd) * 0.21):.2f}\n")
        tot_run = totals_p.add_run(f"TOTAL USD:      {(float(total_usd) * 1.21):.2f}")
    else:
        tot_run = totals_p.add_run(f"TOTAL USD:      {float(total_usd):.2f}")
    tot_run.bold = True
    tot_run.font.size = Pt(12)

    doc.add_paragraph()

    # 7. SCOPE
    scope_title = doc.add_paragraph()
    scope_title.paragraph_format.keep_with_next = True
    run = scope_title.add_run("SCOPE")
    run.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph().add_run("Includes:\n").bold = True
    doc.paragraphs[-1].add_run(scope_includes)
    
    doc.add_paragraph().add_run("Excludes:\n").bold = True
    doc.paragraphs[-1].add_run(scope_excludes)

    doc.add_paragraph()

    # 8. TERMS
    terms_title = doc.add_paragraph()
    terms_title.paragraph_format.keep_with_next = True
    run = terms_title.add_run("TERMS")
    run.bold = True
    run.font.size = Pt(11)

    terms_data = [
        ("Currency", "USD"),
        ("Offer validity", offer_validity),
        ("Payment terms", payment_terms),
        ("Delivery time", f"{delivery_time} business days" if delivery_time else "[e.g. 5 business days from PO]"),
        ("Place of delivery", f"{puerto} / on board {buque}" if puerto and buque else "[port / warehouse / on board M/V ____]"),
    ]
    
    terms_table = doc.add_table(rows=len(terms_data), cols=2)
    terms_table.autofit = False
    for i, (label, val) in enumerate(terms_data):
        terms_table.cell(i, 0).width = Cm(4.0)
        p_label = terms_table.cell(i, 0).paragraphs[0]
        p_label.add_run(label).bold = True
        
        terms_table.cell(i, 1).width = Cm(14.0)
        terms_table.cell(i, 1).text = val

    doc.add_paragraph()

    # 9. NOTES
    notes_title = doc.add_paragraph()
    notes_title.paragraph_format.keep_with_next = True
    run = notes_title.add_run("NOTES")
    run.bold = True
    run.font.size = Pt(11)
    
    notes_lines = [
        "Prices subject to confirmation of stock and availability at the time of the purchase order.",
        "Quantities and specifications to be confirmed by the client before dispatch."
    ]
    if notes and str(notes).strip() != "":
        notes_lines.append(str(notes).strip())

    for note in notes_lines:
        p = doc.add_paragraph(f"– {note}")
        p.paragraph_format.left_indent = Cm(0.5)

    # 10. FIRMA
    sig_p = doc.add_paragraph("Yours faithfully,\n\n")
    sig_name = sig_p.add_run("Eva Proios\n")
    sig_name.bold = True
    sig_p.add_run("Operations – Proios S.A.\neva@proios.com · +549 11 57265031")
    sig_p.paragraph_format.space_after = Pt(20)

    # 11. FOOTER
    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run("Proios S.A. · Comodoro Pedro Zanni 351 floor 5th 503 LN. Buenos Aires (C1104AAH) Argentina | +549 11 57265031 | eva@proios.com | WWW.PROIOS.COM")
    f_run.font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    
    return docx_bytes

def generar_cotizacion_pdf(operacion, offer_validity="15 days", payment_terms="30 days from invoice date", delivery_time="5", include_vat=True, scope_includes="[detail what the supply / service comprises]", scope_excludes="[freight, customs clearance, additional labour, parts not listed, etc.]", notes="[Other relevant note]", attn="Operations / Technical Department"):
    import tempfile
    import os
    import subprocess
    
    # 1. Generar DOCX con todo el estilo y estructura
    docx_bytes = generar_cotizacion_docx(operacion, offer_validity, payment_terms, delivery_time, include_vat, scope_includes, scope_excludes, notes, attn)
    
    # 2. Guardar a disco
    fd_docx, temp_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd_docx)
    with open(temp_docx, 'wb') as f:
        f.write(docx_bytes)
        
    temp_dir = os.path.dirname(temp_docx)
    base_name = os.path.splitext(os.path.basename(temp_docx))[0]
    temp_pdf = os.path.join(temp_dir, f"{base_name}.pdf")
    
    # 3. Convertir usando LibreOffice en background
    try:
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', temp_docx, '--outdir', temp_dir],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        
        with open(temp_pdf, 'rb') as f:
            pdf_bytes = f.read()
    finally:
        if os.path.exists(temp_docx): os.remove(temp_docx)
        if os.path.exists(temp_pdf): os.remove(temp_pdf)
        
    return pdf_bytes



def generar_cotizacion_servicio_docx(operacion, user, params):
    # Generar desde un docx en blanco
    doc = Document()
    
    # Configurar márgenes si empezamos de cero (tamaño A4 estándar)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # Configurar logo en el header usando la función helper que ya tienes
    insert_header_logo(doc)
    
    # Simular footer si empezamos de cero
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Proios S.A. · Comodoro Pedro Zanni 351 floor 5th 503 LN. Buenos Aires (C1104AAH) Argentina | +549 11 57265031 | eva@proios.com | WWW.PROIOS.COM")
    run.font.color.rgb = RGBColor(0, 0, 255) # Azul
    run.font.size = Pt(8)
    run.font.name = 'Arial'

    # Modificar el estilo "Normal" fue removido porque python-docx corrompe
    # el XML base si se altera directamente el style.font.name sin configurar
    # las fuentes eastAsia y complexScript correspondientes.
    # Usaremos el helper add_run_arial para todas las inserciones.

    # ------------------ HELPERS ------------------
    def add_run_arial(paragraph, text, bold=False, italic=False, underline=False, color=None):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        if color:
            run.font.color.rgb = color
        run.font.name = 'Arial'
        return run

    def add_text_with_red_asterisks(paragraph, text):
        """Helper para resaltar (*) y (**) en color rojo (RGB: 255, 0, 0)"""
        parts = re.split(r'(\(\*\*\)|\(\*\))', text)
        for part in parts:
            if part in ['(*)', '(**)']:
                add_run_arial(paragraph, part, bold=True, color=RGBColor(255, 0, 0))
            else:
                add_run_arial(paragraph, part)

    # ------------------ CONTENIDO ------------------
    
    # 1. Saludo
    p_greeting = doc.add_paragraph()
    cliente_nombre = operacion.cliente.name if operacion.cliente else "Client"
    add_run_arial(p_greeting, f"Dear Ms./Mrs. {cliente_nombre},")
    
    doc.add_paragraph() # Espaciado
    
    # 2. Detalle del servicio
    p_intro = doc.add_paragraph()
    intro_text = operacion.detalle_servicio or "Please note below our quotation for the requested services."
    add_run_arial(p_intro, intro_text)
    
    doc.add_paragraph() # Espaciado
    
    # 3. Transporte (antes de los items, según el screenshot)
    if params.get('transport'):
        p_transport = doc.add_paragraph()
        add_run_arial(p_transport, params['transport'])
        doc.add_paragraph()
    
    # 4. Items (Numeración y listado)
    p_items_title = doc.add_paragraph()
    add_run_arial(p_items_title, "1. Items:", bold=True)
    
    # Generar letras A, B, C... para los sub-items
    letters = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    
    custom_items_str = params.get('custom_items')
    if custom_items_str:
        try:
            items_to_render = json.loads(custom_items_str)
        except:
            items_to_render = []
    else:
        # Fallback al original si existieran en DB
        from apps.inventario.models import Articulo
        items_to_render = []
        for det in operacion.detalles.all():
            try:
                articulo = Articulo.objects.get(id=det.articulo_id)
                nombre = articulo.nombre
            except Articulo.DoesNotExist:
                nombre = f"Item #{det.articulo_id}"
            items_to_render.append({
                'nombre': nombre,
                'cantidad': det.cantidad,
                'precio_unitario': det.precio_unitario
            })
        
    for idx, item in enumerate(items_to_render):
        p_item = doc.add_paragraph()
        p_item.paragraph_format.left_indent = Cm(0.63) # Indentación para las letras
        
        letter = letters[idx] if idx < len(letters) else str(idx)
        add_run_arial(p_item, f"{letter}. ", bold=True)
        
        # El nombre del artículo y los asteriscos
        nombre_item = item.get('nombre', 'Item')
        cantidad = float(item.get('cantidad', 1) or 1)
        precio_unitario = float(item.get('precio_unitario', 0) or 0)
        
        add_text_with_red_asterisks(p_item, f"{nombre_item}: ")
        
        # Cálculo del precio, diferenciando si es tarifa diaria o valor total
        if "/day" in nombre_item.lower():
            add_run_arial(p_item, f"USD {precio_unitario:,.2f}/day", bold=True)
        else:
            total = precio_unitario * cantidad
            add_run_arial(p_item, f"USD {total:,.2f}", bold=True)
        
        add_run_arial(p_item, " + taxes.")

    doc.add_paragraph()
    
    # 5. Tiempos (mobilization / execution)
    if params.get('mobilization'):
        p_mob = doc.add_paragraph()
        add_run_arial(p_mob, "Estimated time for mobilization: ", italic=True, underline=True)
        add_run_arial(p_mob, f"{params.get('mobilization')}", italic=True, underline=True)
        
    if params.get('execution'):
        p_exec = doc.add_paragraph()
        add_run_arial(p_exec, "Estimated time for the execution: ", italic=True, underline=True)
        add_run_arial(p_exec, f"{params.get('execution')}", italic=True, underline=True)
    
    # Agregar espacio solo si mostramos algún tiempo
    if params.get('mobilization') or params.get('execution'):
        doc.add_paragraph()

    # 6. Texto de cotización adicional (Asteriscos, viñetas, notas rojas)
    if operacion.texto_cotizacion_adicional:
        for line in operacion.texto_cotizacion_adicional.split('\n'):
            if line.strip():
                # Hacemos un salto de página justo antes del texto que pediste para forzar la hoja 2
                if "In case the execution of services during overtime" in line:
                    doc.add_page_break()
                    
                p_add = doc.add_paragraph()
                add_text_with_red_asterisks(p_add, line.strip())

        doc.add_paragraph()

    # 7. Administrative
    p_admin = doc.add_paragraph()
    add_run_arial(p_admin, "2. Administrative:", bold=True)
    
    admin_bullets = []
    if params.get('bank_charges'): admin_bullets.append(f"Bank charges: {params['bank_charges']}")
    if params.get('taxes'): admin_bullets.append(f"Taxes {params['taxes']}")
    if params.get('payment_terms'): admin_bullets.append(f"Payment terms: {params['payment_terms']}")

    for bullet in admin_bullets:
        p_b = doc.add_paragraph()
        p_b.paragraph_format.left_indent = Cm(1.0)
        # Símbolo de viñeta (bullet) manual
        add_run_arial(p_b, f"• {bullet}")
        
    doc.add_paragraph()

    # 8. Remarks
    p_rem = doc.add_paragraph()
    add_run_arial(p_rem, "3. Remarks:", bold=True)
    
    remarks = [
        "The final invoice may vary depending on actual service duration, onboard conditions, and any extra time or resources required.",
        params.get('transport', '')
    ]
    for rem in remarks:
        if rem.strip():
            p_r = doc.add_paragraph()
            p_r.paragraph_format.left_indent = Cm(1.0)
            add_run_arial(p_r, "• ")
            # Las notas de remarks en el PDF original están en negrita y subrayadas
            add_run_arial(p_r, rem.strip(), bold=True, underline=True)

    # 9. Cierre
    p_check = doc.add_paragraph()
    add_run_arial(p_check, "Please check and advise.")

    # 10. Firma (Datos del User)
    p_sig = doc.add_paragraph()
    add_run_arial(p_sig, "Best regards,\n")
    
    # Para maquetar similar a la firma del screenshot (logo al lado del texto), 
    # la opción más robusta por código es usar una tabla de 1 fila, 2 columnas, sin bordes.
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    
    # Columna Izquierda (Espacio para el logo de la firma si hubiese)
    cell_logo = sig_table.cell(0, 0)
    cell_logo.width = Cm(5.0)
    # cell_logo.paragraphs[0].add_run().add_picture('metalock_signature_logo.png', width=Cm(4.0))
    
    # Columna Derecha (Datos del User)
    cell_data = sig_table.cell(0, 1)
    cell_data.width = Cm(10.0)
    p_data = cell_data.paragraphs[0]
    
    user_first_name = user.first_name if user and user.first_name else ""
    user_last_name = user.last_name if user and user.last_name else ""
    user_name = f"{user_first_name} {user_last_name}".strip() or (user.username if user else "Proios Representative")
    
    add_run_arial(p_data, f"{user_name}\n", bold=True)
    add_run_arial(p_data, "Comercial\n", bold=True)
    add_run_arial(p_data, "Tel.: +55 13 3226-4686\n") # Teléfono fijo de la empresa
    add_run_arial(p_data, "Cel.: +55 13 99697-3059\n") # Celular (podría venir de user en el futuro)
    
    user_email = "operations@proios.com"
    if user_email:
        run_email = add_run_arial(p_data, f"{user_email}")
        run_email.font.color.rgb = RGBColor(0, 0, 255)
        run_email.underline = True
    
    # Word requires the document to end with a paragraph, not a table.
    # We add an empty paragraph to avoid the "Unreadable content" error.
    last_p = doc.add_paragraph()
    last_p.paragraph_format.space_after = Pt(0)
    last_p.paragraph_format.space_before = Pt(0)
    last_run = last_p.add_run()
    last_run.font.size = Pt(1)
    
    # ==========================================
    # GUARDAR DOCUMENTO
    # ==========================================
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    
    return docx_bytes

from docxtpl import DocxTemplate

def generar_remito_docx(operacion):
    """
    Genera el remito manteniendo 100% los espacios y formato originales.
    Utiliza docxtpl para renderizar sobre una plantilla sin romper la maquetación.
    """
    # Cargar la plantilla.
    template_path = os.path.join(settings.BASE_DIR, 'static_local', 'REMITO_TEMPLATE.docx')
    if not os.path.exists(template_path):
        template_path = os.path.join(settings.BASE_DIR, 'REMITO_TEMPLATE.docx')
        
    doc = DocxTemplate(template_path)
    
    now = datetime.now()
    dia = now.strftime("%d")
    mes = now.strftime("%m")
    anio = now.strftime("%Y")
    
    buque = operacion.ship.name if operacion.ship else ''
    puerto = operacion.port.name if operacion.port else ''
    buque_y_puerto = f"{buque} / {puerto}"
    
    from apps.inventario.models import Articulo
    
    detalles = list(operacion.detalles.all())
    items = []
    for idx, det in enumerate(detalles, start=1):
        try:
            articulo = Articulo.objects.get(id=det.articulo_id)
            nombre = articulo.nombre
        except Articulo.DoesNotExist:
            nombre = f"Articulo #{det.articulo_id}"
            
        cantidad = det.cantidad
        items.append(f"ITEM {idx} : {nombre} : {cantidad} UNIDADES")
        
    texto_items = "\n\n".join(items)
    
    # Rellenar con saltos de línea si hay menos de 5 ítems para mantener la firma fija
    items_faltantes = 5 - len(detalles)
    if items_faltantes > 0:
        # Por cada ítem faltante, agregamos su espacio (un doble salto de línea)
        texto_items += "\n\n" * items_faltantes
        
    context = {
        'dia': dia,
        'mes': mes,
        'anio': anio,
        'buque_y_puerto': buque_y_puerto,
        'items': texto_items
    }
    
    doc.render(context)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    
    return docx_bytes

def generar_remito_pdf(operacion):
    """
    Genera el remito en DOCX respetando la plantilla original y luego lo convierte
    nativamente a PDF utilizando LibreOffice (headless), que corre en el contenedor Linux.
    Esto garantiza 100% de fidelidad con el diseño de Microsoft Word sin usar COM/Windows.
    """
    import tempfile
    import subprocess
    
    # 1. Obtenemos el DOCX con los datos inyectados en la plantilla
    docx_bytes = generar_remito_docx(operacion)
    
    # 2. Guardamos en disco temporal
    fd_docx, temp_docx = tempfile.mkstemp(suffix=".docx")
    with os.fdopen(fd_docx, 'wb') as f:
        f.write(docx_bytes)
        
    temp_dir = os.path.dirname(temp_docx)
    
    try:
        # 3. Convertir a PDF usando libreoffice
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", temp_dir,
            temp_docx
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        # 4. Leer el PDF generado
        base_name = os.path.splitext(os.path.basename(temp_docx))[0]
        pdf_path = os.path.join(temp_dir, f"{base_name}.pdf")
        
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
            
        # 5. Limpieza del PDF
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
            
        return pdf_bytes
    finally:
        # 6. Limpieza del DOCX
        if os.path.exists(temp_docx):
            os.remove(temp_docx)

